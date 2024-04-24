import tkinter as tk
from module import mass
from module import plankla
from module import svet
def krugmihail(canvas, x, y, r, color):
    return canvas.create_oval(x-r, y-r, x+r, y+r, fill=color)

root = tk.Tk()

canvas = tk.Canvas(root, width=1000, height=400)
canvas.pack()

r = 50
electronchik = krugmihail(canvas, 100, 100, r, "red")  
canvas.create_text(100, 100, text="-", font=("Arial", 30), fill="white")

protonchik = krugmihail(canvas, 400, 100, r, "green")  
canvas.create_text(400, 100, text="+", font=("Arial", 30), fill="white")

neitronchik = krugmihail(canvas, 300, 200, r, "blue")  
canvas.create_text(300, 200, text="", font=("Arial", 30), fill="white")
## высчитывание
def udele():
    zar = float(zar_entry.get())
    m = float(m_entry.get())
    puk = zar / m
    result_label.config(text=f"Удельный заряд: {puk}")

root = tk.Tk()
root.title("Расчет удельного заряда")

zar_label = tk.Label(root, text="Заряд электрона:")
zar_label.pack()

zar_entry = tk.Entry(root)
zar_entry.pack()

m_label = tk.Label(root, text="Масса электрона:")
m_label.pack()

m_entry = tk.Entry(root)
m_entry.pack()

shetik_button = tk.Button(root, text="Посчитать", command=udele)
shetik_button.pack()

result_label = tk.Label(root, text="")
result_label.pack()

root.mainloop()
udele = (mass.ze)/(mass.me)
udelp = (mass.zp)/(mass.mp)
print(f'Удельный заряд электрона {udele}')
print(f'Удельный заряд протона {udelp}') 
volnae= (plankla.h)/(svet.c*mass.me)
volnap=(plankla.h)/(svet.c*mass.mp)
volnan=(plankla.h)/(svet.c*mass.mn)
print(f'Комптоновская длина волны электрона {volnae}')
print(f'Комптоновская длина волны протона {volnap}')
print(f'Комптоновская длина волны нейтрона {volnan}')
from docx import Document
doc = Document()
doc.add_heading('Отчет', level=1)
doc.add_heading('Электроны,протоны,нейтроны', level = 2)
doc.add_paragraph('Импортируем библиотеку, создаем холст')
doc.add_paragraph('Электрон создается с x,y = 100,100 цвета красного с минусом внутри')
doc.add_paragraph('Протон создается с x,y = 400,100 цвета зеленого и с плюсиком внутри(положительный заряд)')
doc.add_paragraph('Нейтрон создается с x,y = 300,200 цвета синего.')
doc.add_paragraph('Создаем пакет, состоящий из трех модулей, в которых расчитывается масса всех частиц и постоянные числа.')

doc.add_heading('Удельный заряд', level = 2)
doc.add_paragraph('Создаем функцию со значениями и самим выражением')
doc.add_paragraph('Создаем кнопки для введение чисел заряда и масса')
doc.add_paragraph('Создаем кнопку для счета выражения')
doc.add_paragraph('Выводим результат')

doc.add_heading('Комоптоновская длина волны,удельный заряд', level = 2)
doc.add_paragraph('Импортируем библиотеки')
doc.add_paragraph('Высчитываем комптоновской длину волны,расчет удельного заряда')
doc.add_paragraph('Выводим')
doc.save('Отчет2.docx')